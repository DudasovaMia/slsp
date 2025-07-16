import re
import json
from docx import Document
import os
import requests

def load_docx(path):
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def load_law_docx(path):
    doc = Document(path)
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

def load_law_txt(path):
    with open(path, encoding='utf-8') as f:
        return f.read()

def split_law_to_dict(law_text):
    law_dict = {}
    paragrafy = re.split(r'(§\s*\d+[a-zA-Z]*)', law_text)
    for i in range(1, len(paragrafy), 2):
        paragraf = paragrafy[i].strip()
        text = paragrafy[i+1].strip()
        odseky = re.split(r'\n\s*(\d+)\.', text)
        if len(odseky) > 1:
            for j in range(1, len(odseky), 2):
                odsek = odseky[j].strip()
                odsek_text = odseky[j+1].strip()
                key = f"{paragraf}, ods. {odsek}"
                law_dict[key] = odsek_text
        else:
            law_dict[paragraf] = text
    return law_dict

def build_prompt(law_text, contract_text):
    return f'''
Porovnaj nasledujúcu zmluvnú klauzulu so zákonným ustanovením (odpovedaj ano alebo nie, ak nie doplň dôvod).

Zákonné ustanovenie:
"""{law_text}"""

Zmluvná klauzula:
"""{contract_text}"""

Otázka:
Je zmluvná klauzula v súlade so zákonom? Odpovedz ÁNO alebo NIE a vysvetli prečo.
'''

def call_local_llm(law_text, contract_text):
    prompt = (
        "Si právny expert na kolektívne investovanie. "
        + build_prompt(law_text, contract_text)
    )
    try:
        response = requests.post(
            "http://localhost:1234/v1/chat/completions",
            headers={"Content-Type": "application/json"},
            json={
                "model": "mistral-7b-instruct-v0.2",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.3,
                "max_tokens": 500
            }
        )
        answer = response.json()["choices"][0]["message"]["content"]
        if "nie" in answer.lower():
            return {"sulad": False, "vysvetlenie": answer.strip()}
        else:
            return {"sulad": True, "vysvetlenie": answer.strip()}
    except Exception as e:
        return {"sulad": False, "vysvetlenie": f"Chyba pri volaní modelu: {str(e)}"}

def compare_law_and_contract(law_dict, contract_parts):
    results = []
    filter_paragraphs = ["§3", "§5"]
    for law_key, law_text in law_dict.items():
        if not any(law_key.startswith(f) for f in filter_paragraphs):
            continue
        for idx, contract_part in enumerate(contract_parts):
            print(f"\nPorovnávam: {law_key} vs. časť zmluvy {idx+1}/{len(contract_parts)}")
            result = call_local_llm(law_text, contract_part)
            print(f"Výsledok: {'NESÚLAD' if not result['sulad'] else 'SÚLAD'}")
            print(f"Odpoveď modelu: {result['vysvetlenie']}")
            results.append({
                "zakon": law_key,
                "cast_zmluvy": contract_part,
                "vysvetlenie": result["vysvetlenie"],
                "sulad": result["sulad"]
            })
    print("\nPorovnávanie dokončené.\n")
    return results

def save_results(results, path="vysledky.json"):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

def print_summary(results):
    print("Paragraf zákona\tČasť zmluvy\tSúlad\tVysvetlenie")
    for r in results:
        print(f"{r['zakon']}\t{r['cast_zmluvy'][:40]}...\t{r['sulad']}\t{r['vysvetlenie']}")

if __name__ == "__main__":
    zmluva = "zmluvy/vzor_zmluva_01.docx"
    zakon = "zakony/ZZ_2011_203_20240301.docx"

    paragrafy = load_docx(zmluva)
    print(f"Načítaných častí zmluvy: {len(paragrafy)}")

    law_text = load_law_docx(zakon)
    law_dict = split_law_to_dict(law_text)
    print(f"Načítaných paragrafov/odsekov zákona: {len(law_dict)}")

    results = compare_law_and_contract(law_dict, paragrafy)

    save_results(results)

    print_summary(results)

    print("\nUkážka prvých 5 odpovedí modelu:")
    for r in results[:5]:
        print(json.dumps(r, ensure_ascii=False, indent=2))

